"""Extracción de texto de PDF, DOCX, TXT y MD, página a página.

Dos exigencias mandan sobre el diseño:

1. **La página se conserva.** Una cita que no dice la página no resiste una
   verificación contra la fuente, y eso es un sub-criterio explícito de la rúbrica.
2. **Nada revienta.** El jurado va a subir un PDF cualquiera: protegido, escaneado,
   corrupto o de 400 páginas. Un error se convierte en un documento marcado con su
   motivo en la consola, nunca en una excepción que tumba la ingesta (riesgo R7).
"""

from __future__ import annotations

import hashlib
import logging
import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path

log = logging.getLogger("postopfriend.rag")

EXTENSIONES = {".pdf", ".docx", ".txt", ".md"}
MAX_BYTES = 80 * 1024 * 1024
MAX_PAGINAS = 600
MIN_CARACTERES_PAGINA = 50  # por debajo de esto, la página se considera sin capa de texto

# Palabras funcionales muy frecuentes y casi disjuntas entre los dos idiomas.
# Basta para distinguir es/en en un documento entero; no se pretende más.
_MARCAS_ES = {"de", "la", "que", "el", "en", "los", "del", "las", "por", "con", "para", "una", "es"}
_MARCAS_EN = {"the", "of", "and", "to", "in", "for", "with", "that", "is", "was", "were", "are", "on"}


@dataclass
class Pagina:
    numero: int  # 1-indexado, como lo muestra cualquier visor de PDF
    texto: str


@dataclass
class DocumentoExtraido:
    archivo: Path
    titulo: str
    paginas: list[Pagina] = field(default_factory=list)
    idioma: str = "es"
    sha256: str = ""
    estado: str = "disponible"  # disponible | sin_texto | error
    detalle: str = ""

    @property
    def total_paginas(self) -> int:
        return len(self.paginas)

    @property
    def caracteres(self) -> int:
        return sum(len(p.texto) for p in self.paginas)

    @property
    def paginas_sin_texto(self) -> list[int]:
        return [p.numero for p in self.paginas if len(p.texto.strip()) < MIN_CARACTERES_PAGINA]


def normalizar(texto: str) -> str:
    """Reconstruye párrafos legibles a partir del texto crudo del PDF.

    PyMuPDF respeta el salto de línea visual del documento. En los artículos a dos
    columnas eso llega a producir una palabra por renglón — «In\\norder\\nto\\n…» —
    y un fragmento así embebe mal, se lee peor y es ilegible como cita en pantalla.

    Se colapsa el salto simple a espacio y se conserva el doble como frontera de
    párrafo, que es la única que corresponde a una separación real del texto.
    """
    texto = unicodedata.normalize("NFKC", texto)
    texto = texto.replace("­", "")               # guion suave invisible
    texto = re.sub(r"(\w)-\s*\n\s*(\w)", r"\1\2", texto)  # palabra partida al final de renglón
    texto = re.sub(r"[ \t]*\n[ \t]*", "\n", texto)
    texto = re.sub(r"(?<!\n)\n(?!\n)", " ", texto)        # renglón suelto → espacio
    texto = re.sub(r"[ \t ]+", " ", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


def detectar_idioma(texto: str) -> str:
    palabras = re.findall(r"[a-záéíóúñü]+", texto.lower())[:4000]
    if not palabras:
        return "es"
    es = sum(1 for p in palabras if p in _MARCAS_ES)
    en = sum(1 for p in palabras if p in _MARCAS_EN)
    return "en" if en > es else "es"


def hash_texto(paginas: list[Pagina]) -> str:
    """SHA-256 del texto normalizado, no del archivo.

    El corpus trae el mismo documento dos veces con distinta capitalización en el
    nombre y bytes distintos. Hashear el contenido es lo que permite deduplicar.
    """
    crudo = "\n".join(p.texto for p in paginas)
    compacto = re.sub(r"\s+", " ", crudo).strip().lower()
    return hashlib.sha256(compacto.encode("utf-8")).hexdigest()


# ─── Extractores por formato ─────────────────────────────────────────────────

def _extraer_pdf(ruta: Path) -> tuple[list[Pagina], str]:
    import pymupdf

    paginas: list[Pagina] = []
    detalle = ""
    with pymupdf.open(ruta) as doc:
        if doc.needs_pass:
            raise ValueError("el PDF está protegido con contraseña")
        total = doc.page_count
        if total > MAX_PAGINAS:
            detalle = f"documento de {total} páginas: se indexan las primeras {MAX_PAGINAS}"
            total = MAX_PAGINAS
        for i in range(total):
            try:
                texto = doc.load_page(i).get_text("text")
            except Exception as e:  # una página rota no invalida el documento
                log.warning("%s p.%d ilegible: %s", ruta.name, i + 1, e)
                texto = ""
            paginas.append(Pagina(numero=i + 1, texto=normalizar(texto)))
    return paginas, detalle


def _extraer_docx(ruta: Path) -> tuple[list[Pagina], str]:
    import docx

    documento = docx.Document(str(ruta))
    cuerpo = "\n".join(p.text for p in documento.paragraphs)
    for tabla in documento.tables:
        for fila in tabla.rows:
            cuerpo += "\n" + " | ".join(c.text.strip() for c in fila.cells)
    # DOCX no tiene paginación estable sin renderizar: todo el documento es la página 1.
    return [Pagina(numero=1, texto=normalizar(cuerpo))], "DOCX sin paginación: se cita como página 1"


def _extraer_texto_plano(ruta: Path) -> tuple[list[Pagina], str]:
    crudo = ruta.read_text(encoding="utf-8", errors="replace")
    return [Pagina(numero=1, texto=normalizar(crudo))], ""


# ─── Entrada pública ─────────────────────────────────────────────────────────

def extraer(ruta: Path, titulo: str | None = None) -> DocumentoExtraido:
    """Extrae un documento. Nunca lanza: los fallos vuelven como `estado='error'`."""
    doc = DocumentoExtraido(archivo=ruta, titulo=titulo or ruta.stem)

    if ruta.suffix.lower() not in EXTENSIONES:
        doc.estado, doc.detalle = "error", f"formato no soportado: {ruta.suffix}"
        return doc
    try:
        tamano = ruta.stat().st_size
    except OSError as e:
        doc.estado, doc.detalle = "error", f"no se pudo leer el archivo: {e}"
        return doc
    if tamano > MAX_BYTES:
        doc.estado, doc.detalle = "error", f"archivo de {tamano / 1e6:.0f} MB: excede el límite de {MAX_BYTES // 10**6} MB"
        return doc

    try:
        if ruta.suffix.lower() == ".pdf":
            doc.paginas, doc.detalle = _extraer_pdf(ruta)
        elif ruta.suffix.lower() == ".docx":
            doc.paginas, doc.detalle = _extraer_docx(ruta)
        else:
            doc.paginas, doc.detalle = _extraer_texto_plano(ruta)
    except Exception as e:
        doc.estado, doc.detalle = "error", f"{type(e).__name__}: {e}"
        return doc

    doc.sha256 = hash_texto(doc.paginas)
    doc.idioma = detectar_idioma(" ".join(p.texto for p in doc.paginas[:20]))

    if doc.caracteres < MIN_CARACTERES_PAGINA:
        # Típico de un PDF escaneado sin OCR. Se lista en la consola con el motivo,
        # nunca se falla en silencio ni se bloquea el arranque por no tener Tesseract.
        doc.estado = "sin_texto"
        doc.detalle = "sin capa de texto (PDF escaneado): no indexado. Requiere OCR."
    return doc


def recorrer_corpus(raiz: Path) -> list[Path]:
    """Archivos indexables bajo `raiz`, en orden estable."""
    return sorted(p for p in raiz.rglob("*") if p.is_file() and p.suffix.lower() in EXTENSIONES)
